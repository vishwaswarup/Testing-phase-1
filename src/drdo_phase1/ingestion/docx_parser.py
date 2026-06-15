"""
Experiment 3 — DOCX Text Extraction using python-docx
======================================================

Purpose
-------
Learn how python-docx reads Microsoft Word (.docx) files and
exposes their paragraph-level structure.

Library
-------
python-docx  →  `from docx import Document`

Key Concepts
------------
* `Document(path)` opens the file.
* `doc.paragraphs` is a list of Paragraph objects.
* Each paragraph has a `.text` attribute with its string content.
* Paragraphs include headings, body text, list items, etc.
"""

import os
from docx import Document


def extract_docx(file_path: str) -> dict:
    """
    Open a DOCX file and extract all paragraph text.

    Returns
    -------
    dict with keys:
        text            – full extracted text
        paragraph_count – number of paragraphs
        extractor       – name of the extractor used
    """

    # --- 1. Validate path ---------------------------------------------------
    if not os.path.isfile(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    try:
        doc = Document(file_path)
    except Exception as e:
        raise RuntimeError(f"Could not open DOCX: {e}")

    # --- 2. Read paragraphs -------------------------------------------------
    paragraphs = doc.paragraphs
    paragraph_text = "\n".join(para.text for para in paragraphs)

    # --- 3. Read tables (Phase 1B: doc.tables for tabular data) -------------
    table_texts = []
    for table_idx, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            table_texts.append(f"\n[Table {table_idx + 1}]\n" + "\n".join(rows))

    full_text = paragraph_text
    if table_texts:
        full_text += "\n" + "\n".join(table_texts)

    return {
        "text": full_text,
        "paragraph_count": len(paragraphs),
        "table_count": len(doc.tables),
        "extractor": "python-docx",
    }


def main():
    print("=" * 60)
    print("  Experiment 3 — DOCX Extraction (python-docx)")
    print("=" * 60)

    file_path = input("\nEnter DOCX path: ").strip()

    try:
        result = extract_docx(file_path)
    except (FileNotFoundError, RuntimeError) as e:
        print(f"[ERROR] {e}")
        return

    print("\n--- Paragraphs ---\n")
    for i, line in enumerate(result["text"].splitlines(), start=1):
        print(f"  [{i}] {line}")

    print()
    print("-" * 40)
    print(f"  Paragraph Count : {result['paragraph_count']}")
    print(f"  Character Count : {len(result['text'])}")
    print()


if __name__ == "__main__":
    main()
